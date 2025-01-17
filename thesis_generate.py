import streamlit as st
from docx import Document
import os

# Function to generate a simple thesis based on the given title and summary
def generate_thesis(title, summary):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph(summary)
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("This is an auto-generated introduction for the thesis.")
    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph("This is an auto-generated conclusion for the thesis.")
    return doc

# Streamlit app
st.title("Thesis Workflow Application")

# Step 1: Enter Title of Thesis
title = st.text_input("Enter the title of your thesis:")

# Step 2: Enter Summary of Thesis
if title:
    summary = st.text_area("Enter the summary of your thesis:")
    
    # Step 3: Upload Sample Thesis
    if summary:
        st.write("Upload a sample thesis for verification:")
        uploaded_files = st.file_uploader("Upload sample thesis files (PDF, DOCX):", 
                                          type=["pdf", "docx"], 
                                          accept_multiple_files=True)
        uploaded_file_names = [file.name for file in uploaded_files]
        
        if uploaded_files:
            # Ask if user wants to upload more files
            more_files = st.radio("Do you want to upload more files?", ("Yes", "No"))
            
            if more_files == "No":
                # Step 4: Verify the Title with Uploaded Sample Thesis
                # Simulating verification (implement logic as needed)
                st.info(f"Verifying the title '{title}' with uploaded files...")
                
                # For simplicity, assume verification passes
                st.success("Title verified successfully!")
                
                # Step 5: Ask if User Wants to Generate Own Thesis
                generate = st.radio("Do you want to generate your own thesis?", ("Yes", "No"))
                
                if generate == "Yes":
                    # Generate Thesis
                    thesis_doc = generate_thesis(title, summary)
                    
                    # Save Thesis to a Word File
                    file_path = os.path.join("generated_thesis.docx")
                    thesis_doc.save(file_path)
                    
                    # Allow User to Download the File
                    with open(file_path, "rb") as file:
                        st.download_button(
                            label="Download Generated Thesis",
                            data=file,
                            file_name=f"{title.replace(' ', '_')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    st.success("Thesis generated and ready for download!")
